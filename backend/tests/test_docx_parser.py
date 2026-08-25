from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
import pytest

from app.services.docx_parser import extract_docx_text, validate_docx_file_bytes


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("合同标题")
    document.add_paragraph("  ")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "甲方联系人"
    table.rows[0].cells[1].text = "张三"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_docx_text_includes_paragraphs_and_tables() -> None:
    text = extract_docx_text(_build_docx_bytes())

    assert "合同标题" in text
    assert "甲方联系人 | 张三" in text


def test_docx_validation_rejects_non_docx_zip_archives() -> None:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("notes.txt", "not a Word document")

    with pytest.raises(ValueError, match="word/document.xml"):
        validate_docx_file_bytes(buffer.getvalue())


def test_docx_validation_rejects_invalid_archives() -> None:
    with pytest.raises(ValueError, match="valid ZIP archive"):
        validate_docx_file_bytes(b"not-a-docx")
