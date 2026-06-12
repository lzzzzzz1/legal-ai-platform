import json
from io import BytesIO
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient
from lxml import etree

from app.main import app


client = TestClient(app)
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _build_docx_bytes() -> bytes:
    document = Document()
    styled_paragraph = document.add_paragraph("合同份数：一式两份。")
    styled_paragraph.style = "List Paragraph"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "签订地点：未约定。"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_english_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("(i) Representation by Counsel. Each of the parties acknowledges that it has entered into this Agreement based upon its independent judgment.")
    document.add_paragraph("(j) Counterparts. This Agreement may be executed in one or more counterparts, each of which shall be considered an original instrument.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _read_docx_xml(docx_bytes: bytes, path: str) -> etree._Element:
    with ZipFile(BytesIO(docx_bytes), "r") as archive:
        return etree.fromstring(archive.read(path))


def _paragraph_texts_from_xml(root: etree._Element) -> list[str]:
    paragraphs = []
    for paragraph in root.findall(".//w:p", W_NS):
        parts = [
            node.text or ""
            for node in paragraph.xpath(".//w:t | .//w:delText", namespaces=W_NS)
        ]
        joined = "".join(parts).strip()
        if joined:
            paragraphs.append(joined)
    return paragraphs


def test_export_returns_modified_docx() -> None:
    modifications = [
        {
            "original": "合同份数：一式两份。",
            "modified": "合同份数：一式四份，甲乙双方各执两份。",
        },
        {
            "original": "签订地点：未约定。",
            "modified": "签订地点：上海市浦东新区。",
        },
    ]

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": json.dumps(modifications, ensure_ascii=False)},
    )

    assert response.status_code == 200
    assert response.headers["content-disposition"] == 'attachment; filename="reviewed_contract.docx"'

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    settings_xml = _read_docx_xml(response.content, "word/settings.xml")

    assert document_xml.find(".//w:ins", W_NS) is not None
    assert document_xml.find(".//w:del", W_NS) is not None
    assert settings_xml.find(".//w:trackRevisions", W_NS) is not None


def test_export_fuzzy_matches_replacement_text() -> None:
    modifications = [
        {
            "original": "合同份数一式两份",
            "modified": "合同份数：一式三份，甲乙双方各执一份，存档一份。",
        }
    ]

    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": json.dumps(modifications, ensure_ascii=False)},
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")

    assert document_xml.find(".//w:ins", W_NS) is not None
    assert "合同份数：一式三份，甲乙双方各执一份，存档一份。" in "".join(document_xml.xpath(".//w:t/text()", namespaces=W_NS))


def test_export_inserts_missing_clause_after_anchor() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "新增税务条款：税费由乙方承担。",
                        "insert_after_text": "合同份数：一式两份。",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs = _paragraph_texts_from_xml(document_xml)

    assert paragraphs[0] == "合同份数：一式两份。"
    assert "新增税务条款：税费由乙方承担。" in paragraphs[1]
    assert document_xml.find(".//w:ins", W_NS) is not None


def test_export_fuzzy_matches_insertion_anchor() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "新增通知条款：双方应明确联系人与送达邮箱。",
                        "insert_after_text": "合同份数一式两份",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs = _paragraph_texts_from_xml(document_xml)

    assert "新增通知条款：双方应明确联系人与送达邮箱。" in paragraphs[1]


def test_export_appends_missing_clause_modification() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [{"original": "【缺失该约定】", "modified": "新增条款：双方应明确通知联系人。"}],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs_text = "\n".join(_paragraph_texts_from_xml(document_xml))

    assert "新增条款：双方应明确通知联系人。" in paragraphs_text


def test_export_matches_anchor_by_clause_heading() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_english_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "【缺失该约定】",
                        "modified": "Counterparts. This Agreement may be executed in two counterparts.",
                        "insert_after_text": "Counterparts",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs = _paragraph_texts_from_xml(document_xml)

    assert paragraphs[1].startswith("(j) Counterparts.")
    assert paragraphs[2] == "Counterparts. This Agreement may be executed in two counterparts."


def test_export_replaces_text_by_clause_heading_similarity() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_english_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "modifications": json.dumps(
                [
                    {
                        "original": "Counterparts. This Agreement may be executed in one or more counterparts",
                        "modified": "Counterparts. This Agreement may be executed in two counterparts.",
                    }
                ],
                ensure_ascii=False,
            )
        },
    )

    assert response.status_code == 200

    document_xml = _read_docx_xml(response.content, "word/document.xml")
    paragraphs_text = "\n".join(_paragraph_texts_from_xml(document_xml))

    assert "Counterparts. This Agreement may be executed in two counterparts." in paragraphs_text


def test_export_rejects_invalid_modifications_json() -> None:
    response = client.post(
        "/api/export",
        files={
            "file": (
                "contract.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"modifications": "{}"},
    )

    assert response.status_code == 400
