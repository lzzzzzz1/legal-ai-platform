from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document


MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_ARCHIVE_ENTRIES = 5_000
MAX_DOCX_COMPRESSION_RATIO = 120


def validate_docx_file_bytes(file_bytes: bytes) -> None:
    """Reject malformed or disproportionate DOCX archives before parsing.

    DOCX files are ZIP archives. A small uploaded archive can otherwise expand
    to an unexpectedly large payload while python-docx reads it.
    """
    try:
        with ZipFile(BytesIO(file_bytes), "r") as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ValueError("DOCX contains too many archive entries.")
            if "word/document.xml" not in archive.namelist():
                raise ValueError("DOCX is missing word/document.xml.")

            total_uncompressed = 0
            for entry in entries:
                if entry.is_dir():
                    continue
                total_uncompressed += entry.file_size
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ValueError("DOCX expands beyond the allowed size.")
                if entry.compress_size and entry.file_size / entry.compress_size > MAX_DOCX_COMPRESSION_RATIO:
                    raise ValueError("DOCX contains a suspiciously compressed entry.")
    except BadZipFile as exc:
        raise ValueError("Uploaded .docx file is not a valid ZIP archive.") from exc


def extract_docx_text(file_bytes: bytes) -> str:
    validate_docx_file_bytes(file_bytes)
    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    tables = []

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))

    parts = [text for text in paragraphs + tables if text]
    return "\n".join(parts)
