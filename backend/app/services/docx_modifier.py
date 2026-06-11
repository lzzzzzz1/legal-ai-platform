import json
from io import BytesIO
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


Modification = dict[str, Any]


def parse_modifications(modifications_json: str) -> list[Modification]:
    payload = json.loads(modifications_json)
    if not isinstance(payload, list):
        raise ValueError("modifications must be a JSON array")

    for item in payload:
        if not isinstance(item, dict):
            raise ValueError("each modification must be an object")

    return payload


def _modification_texts(modification: Modification) -> tuple[str, str]:
    original = modification.get("original")
    if original is None:
        original = modification.get("original_text")

    modified = modification.get("modified")
    if modified is None:
        modified = modification.get("suggestion")

    if not isinstance(original, str) or not original:
        raise ValueError("each modification requires a non-empty original text")

    if not isinstance(modified, str) or not modified:
        raise ValueError("each modification requires a non-empty modified text")

    return original, modified


def _replace_in_paragraph(paragraph: Paragraph, original: str, modified: str) -> bool:
    paragraph_text = paragraph.text
    if original not in paragraph_text:
        return False

    replaced_text = paragraph_text.replace(original, modified)
    if paragraph.runs:
        paragraph.runs[0].text = replaced_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced_text)

    return True


def _iter_table_paragraphs(document: Document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs


MISSING_SENTINEL = "\u3010\u7f3a\u5931\u8be5\u7ea6\u5b9a\u3011"


def _is_missing_sentinel(text: str) -> bool:
    return text.strip() in (MISSING_SENTINEL, "\u7f3a\u5931\u8be5\u7ea6\u5b9a")


def _clear_paragraph(paragraph: Paragraph) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = ""
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run("")


def replace_docx_body_text(file_bytes: bytes, final_text: str) -> bytes:
    document = Document(BytesIO(file_bytes))
    lines = [line.strip() for line in final_text.splitlines() if line.strip()]

    if not lines:
        raise ValueError("final_text must include readable text")

    if not document.paragraphs:
        document.add_paragraph("")

    target_index = 0

    for line in lines:
        if target_index < len(document.paragraphs):
            paragraph = document.paragraphs[target_index]
            if paragraph.runs:
                paragraph.runs[0].text = line
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(line)
        else:
            document.add_paragraph(line)
        target_index += 1

    for paragraph in document.paragraphs[target_index:]:
        _clear_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def modify_docx_inplace(file_bytes: bytes, modifications: list[Modification]) -> bytes:
    document = Document(BytesIO(file_bytes))
    normalized_modifications = [_modification_texts(item) for item in modifications]

    paragraphs = list(document.paragraphs)
    paragraphs.extend(_iter_table_paragraphs(document))

    for original, modified in normalized_modifications:
        if _is_missing_sentinel(original):
            # Append as new paragraph at end of document
            document.add_paragraph(modified)
        else:
            for paragraph in paragraphs:
                _replace_in_paragraph(paragraph, original, modified)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
